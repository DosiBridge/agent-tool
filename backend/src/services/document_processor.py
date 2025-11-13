"""
Document processing service for file uploads and text extraction
"""
import os
import uuid
import json
from pathlib import Path
from typing import List, Dict, Optional, Tuple
import chardet

# Try to import document processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pypdf as PyPDF2
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False
        print("⚠️  PyPDF2/pypdf not available. PDF processing will be limited.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not available. DOCX processing will be limited.")

from src.core import Config


class DocumentProcessor:
    """Service for processing uploaded documents"""
    
    # Supported file types
    SUPPORTED_TYPES = {
        'pdf': 'application/pdf',
        'txt': 'text/plain',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'md': 'text/markdown',
    }
    
    # Maximum file size (100MB)
    MAX_FILE_SIZE = 100 * 1024 * 1024
    
    def __init__(self, upload_dir: Optional[str] = None):
        """Initialize document processor"""
        if upload_dir is None:
            upload_dir = os.path.join(Config.ROOT_DIR, "uploads")
        
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported"""
        ext = Path(filename).suffix.lower().lstrip('.')
        return ext in self.SUPPORTED_TYPES
    
    def validate_file(self, file_content: bytes, filename: str) -> Tuple[bool, Optional[str]]:
        """Validate uploaded file"""
        # Check file size
        if len(file_content) > self.MAX_FILE_SIZE:
            return False, f"File size exceeds maximum of {self.MAX_FILE_SIZE / (1024*1024):.1f}MB"
        
        # Check file type
        if not self.is_supported_file(filename):
            ext = Path(filename).suffix.lower().lstrip('.')
            return False, f"Unsupported file type: {ext}. Supported types: {', '.join(self.SUPPORTED_TYPES.keys())}"
        
        return True, None
    
    def save_file(self, file_content: bytes, original_filename: str, user_id: int) -> Tuple[str, str]:
        """Save uploaded file to disk"""
        # Create user-specific directory
        user_dir = self.upload_dir / str(user_id)
        user_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate unique filename
        ext = Path(original_filename).suffix
        unique_filename = f"{uuid.uuid4()}{ext}"
        file_path = user_dir / unique_filename
        
        # Save file
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        return str(file_path), unique_filename
    
    def extract_text(self, file_path: str, file_type: str) -> Tuple[str, Dict]:
        """Extract text from document"""
        file_path_obj = Path(file_path)
        
        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        metadata = {
            "file_type": file_type,
            "file_size": file_path_obj.stat().st_size,
        }
        
        if file_type == 'pdf':
            return self._extract_from_pdf(file_path, metadata)
        elif file_type == 'txt' or file_type == 'md':
            return self._extract_from_text(file_path, metadata)
        elif file_type == 'docx':
            return self._extract_from_docx(file_path, metadata)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_from_pdf(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        """Extract text from PDF"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 is required for PDF processing")
        
        text_parts = []
        page_count = 0
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_parts.append(page_text)
                            metadata[f"page_{page_num}"] = len(page_text)
                    except Exception as e:
                        print(f"⚠️  Error extracting text from page {page_num}: {e}")
                        continue
        except Exception as e:
            raise ValueError(f"Failed to read PDF: {str(e)}")
        
        metadata["page_count"] = page_count
        full_text = "\n\n".join(text_parts)
        
        return full_text, metadata
    
    def _extract_from_text(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        """Extract text from plain text file"""
        # Try to detect encoding
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            detected = chardet.detect(raw_data)
            encoding = detected.get('encoding', 'utf-8')
        
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
        except UnicodeDecodeError:
            # Fallback to utf-8 with error handling
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
        
        metadata["encoding"] = encoding
        metadata["line_count"] = len(text.splitlines())
        
        return text, metadata
    
    def _extract_from_docx(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX processing")
        
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            
            metadata["paragraph_count"] = len(paragraphs)
            
            # Extract tables if any
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_texts.append(row_text)
            
            if table_texts:
                text += "\n\n" + "\n".join(table_texts)
                metadata["table_count"] = len(doc.tables)
            
            return text, metadata
        except Exception as e:
            raise ValueError(f"Failed to read DOCX: {str(e)}")
    
    def chunk_text(self, text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Dict[str, any]]:
        """
        Split text into chunks with overlap
        
        Args:
            text: Text to chunk
            chunk_size: Target size of each chunk (in characters)
            chunk_overlap: Overlap between chunks (in characters)
        
        Returns:
            List of chunk dictionaries with content and metadata
        """
        if not text or not text.strip():
            return []
        
        chunks = []
        start = 0
        chunk_index = 0
        
        # Split by paragraphs first for better semantic boundaries
        paragraphs = text.split('\n\n')
        current_chunk = []
        current_length = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            para_length = len(para)
            
            # If paragraph itself is larger than chunk_size, split it
            if para_length > chunk_size:
                # Save current chunk if any
                if current_chunk:
                    chunk_text = '\n\n'.join(current_chunk)
                    chunks.append({
                        "content": chunk_text,
                        "chunk_index": chunk_index,
                        "metadata": {
                            "start_char": start,
                            "end_char": start + len(chunk_text),
                            "length": len(chunk_text)
                        }
                    })
                    start += len(chunk_text) - chunk_overlap
                    chunk_index += 1
                    current_chunk = []
                    current_length = 0
                
                # Split large paragraph
                words = para.split()
                temp_chunk = []
                temp_length = 0
                
                for word in words:
                    word_with_space = word + ' '
                    word_length = len(word_with_space)
                    
                    if temp_length + word_length > chunk_size and temp_chunk:
                        chunk_text = ' '.join(temp_chunk)
                        chunks.append({
                            "content": chunk_text,
                            "chunk_index": chunk_index,
                            "metadata": {
                                "start_char": start,
                                "end_char": start + len(chunk_text),
                                "length": len(chunk_text)
                            }
                        })
                        start += len(chunk_text) - chunk_overlap
                        chunk_index += 1
                        temp_chunk = []
                        temp_length = 0
                    
                    temp_chunk.append(word)
                    temp_length += word_length
                
                if temp_chunk:
                    current_chunk = temp_chunk
                    current_length = temp_length
            else:
                # Check if adding this paragraph would exceed chunk_size
                if current_length + para_length + 2 > chunk_size and current_chunk:
                    # Save current chunk
                    chunk_text = '\n\n'.join(current_chunk)
                    chunks.append({
                        "content": chunk_text,
                        "chunk_index": chunk_index,
                        "metadata": {
                            "start_char": start,
                            "end_char": start + len(chunk_text),
                            "length": len(chunk_text)
                        }
                    })
                    # Start new chunk with overlap
                    overlap_text = chunk_text[-chunk_overlap:] if len(chunk_text) > chunk_overlap else chunk_text
                    start += len(chunk_text) - len(overlap_text)
                    current_chunk = [overlap_text] if overlap_text else []
                    current_length = len(overlap_text)
                    chunk_index += 1
                
                current_chunk.append(para)
                current_length += para_length + 2  # +2 for '\n\n'
        
        # Add remaining chunk
        if current_chunk:
            chunk_text = '\n\n'.join(current_chunk)
            chunks.append({
                "content": chunk_text,
                "chunk_index": chunk_index,
                "metadata": {
                    "start_char": start,
                    "end_char": start + len(chunk_text),
                    "length": len(chunk_text)
                }
            })
        
        return chunks


# Global instance
document_processor = DocumentProcessor()

